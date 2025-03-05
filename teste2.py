import fitz  # PyMuPDF
import docx
import os
from docx.shared import Inches
from PIL import Image

def replace_text_with_images(pdf_path, docx_path, remove_header=True):
    # Configurações
    texts_to_remove = [
        "A avaliação ergonômica preliminar",
        "De acordo com a NR-1, é responsabilidade"
    ]

    # Abrir documentos
    doc_pdf = fitz.open(pdf_path)
    doc_word = docx.Document(docx_path)

    # 1. Localizar e limpar parágrafos alvo
    target_paragraphs = []
    for i, para in enumerate(doc_word.paragraphs):
        for text in texts_to_remove:
            if text in para.text:
                para.clear()  # Apaga o conteúdo do parágrafo
                target_paragraphs.append((i, para))  # Armazena índice e parágrafo

    # Se não encontrar parágrafos para substituir, sair
    if not target_paragraphs:
        print("⚠ Nenhum parágrafo encontrado para substituição!")
        return

    # 2. Processar cada página do PDF
    temp_images = []
    for page_num in range(len(doc_pdf)):
        page = doc_pdf[page_num]
        
        # Renderizar página como imagem com DPI otimizado
        pix = page.get_pixmap(dpi=150)  
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

        # Reduzir consumo de memória convertendo para RGB
        img = img.convert("RGB")

        # Recorte personalizado (ajuste conforme necessidade)
        if remove_header:
            header_height = int(img.height * 0.05)  # 10% para cabeçalho
            footer_height = int(img.height * 0.05)  # 5% para rodapé
            img = img.crop((0, header_height, img.width, img.height - footer_height))

        # Redimensionar mantendo proporção de forma otimizada
        target_width = 500  # Largura em pixels para evitar conversão de Inches
        img.thumbnail((target_width, img.height), Image.LANCZOS)

        # Salvar imagem temporária
        img_path = f"temp_pdf_page_{page_num}.png"
        img.save(img_path, format="PNG", optimize=True)  # Otimizar PNG
        temp_images.append(img_path)

    # 3. Substituir parágrafos limpos pelas imagens
    index_to_insert = target_paragraphs[0][0]  # Índice do primeiro parágrafo removido

    for img_path in temp_images:
        if index_to_insert < len(doc_word.paragraphs):
            new_para = doc_word.paragraphs[index_to_insert]  # Usa o mesmo índice
        else:
            new_para = doc_word.add_paragraph()  # Se não houver mais parágrafos, cria um novo
        
        run = new_para.add_run()
        run.add_picture(img_path, width=Inches(6.5))

        index_to_insert += 1  # Avança para o próximo local de inserção

    # 4. Limpeza final
    for img_path in temp_images:
        if os.path.exists(img_path):
            os.remove(img_path)

    # Salvar documento final
    output_path = docx_path.replace(".docx", "_CORRIGIDO.docx")
    doc_word.save(output_path)
    print(f"✅ Documento processado com sucesso: {output_path}")

# Uso
replace_text_with_images(
    pdf_path=r"C:\Users\Gabriel\Downloads\teste-pgr\files\27.11.2024 - AEP - T.J. ALIMENTOS LTDA - COZINHA GERAL.pdf",
    docx_path=r"C:\Users\Gabriel\Downloads\teste-pgr\files\teste-novo\NOVO MODELO PGR.docx"
)
