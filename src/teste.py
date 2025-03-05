import fitz  # PyMuPDF
from PIL import Image
from docx import Document
from docx.shared import Inches

def replace_text_with_images(pdf_path, docx_path, remove_header=True):
    doc_pdf = fitz.open(pdf_path)
    doc_word = Document(docx_path)

    texts_to_remove = ["AVALIACAO_ERGONOMICA_PRELIMINAR"]
    paragraphs_to_replace = []

    # Identificar e remover textos específicos do Word
    for para in doc_word.paragraphs:
        for text in texts_to_remove:
            if any(text in para.text for text in texts_to_remove):
                para.clear() 
                paragraphs_to_replace.append(para) 

    # Processar cada página do PDF
    images = []
    for page_number in range(len(doc_pdf)):
        page = doc_pdf[page_number]
        pix = page.get_pixmap(dpi=300)  # Renderizar página como imagem (alta resolução)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

        # Definir altura do cabeçalho a ser removido (5%)
        header_height = int(img.height * 0.05)

        if remove_header:
            img = img.crop((0, header_height, img.width, img.height))  # Remover cabeçalho

        max_width = 5.5  # Definindo uma largura menor para não ocupar tanto espaço
        aspect_ratio = img.height / img.width  # Mantém a proporção da imagem
        max_height = max_width * aspect_ratio * 0.3  # Reduzindo um pouco a altura
        
        # Salvar a imagem temporária
        img_path = f"temp_page_{page_number + 1}.png"
        img.save(img_path, "PNG")
        images.append((img_path, max_width, max_height))
            
    for img_path, width, height in images:
        run = paragraphs_to_replace[1].add_run()
        run.add_picture(img_path, width=Inches(width)) 

    output_path = docx_path.replace(".docx", "_editado.docx")
    doc_word.save(output_path)
    
    print("Texto removido e imagens adicionadas nos locais corretos do Word.")

if __name__ == "__main__":
    # Exemplo de uso:
    pdf_path = r"C:\Users\Gabriel\Downloads\teste-pgr\files\27.11.2024 - AEP - T.J. ALIMENTOS LTDA - COZINHA GERAL.pdf"
    docx_path = r"C:\Users\Gabriel\Downloads\teste-pgr\files\teste-novo\NOVO MODELO PGR - Copia.docx"
    
    replace_text_with_images(pdf_path, docx_path)
