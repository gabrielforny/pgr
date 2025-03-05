
from datetime import datetime
# BIBLIOTECAS
import pandas as pd
import pypandoc, re
from io import StringIO
from datetime import datetime
import locale, time

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import locale
locale.setlocale(locale.LC_ALL, 'pt_BR.utf8')
import psutil
import subprocess
import requests

# from ahk import AHK 
# ahk = AHK()

import pygetwindow as gw

import win32com.client as win32
import pyautogui as p

def kill_process_word():
    process_name = "WINWORD"
    # Verificar se o processo está em execução
    for proc in psutil.process_iter(['name']):
        if process_name.lower() in proc.info['name'].lower():
            # Encerrar o processo
            pid = proc.pid
            subprocess.run(['taskkill', '/PID', str(pid), '/F'])


def paste_content_to_new_document(file_docx_result, search_text):
    # Abrir o Word
    time.sleep(3)
    word = win32.Dispatch('Word.Application')
    word.Visible = True
    
    # Abrir o documento alvo
    doc = word.Documents.Open(file_docx_result)
    time.sleep(1)
    
    # Esperar pela janela com o título parcial "result"
    cont = 0
    janela = ahk.find_window_by_title("PGR", title_match_mode=2)
    while janela is None:
        janela = ahk.find_window_by_title("PGR", title_match_mode=2)
        cont +=1
        time.sleep(1)
        if cont == 60: raise Exception("Timout esperando a janela do arquivo result PGR")

    # Localizar o texto "INVENTÁRIO DE RISCOS."
    time.sleep(1)
    find = word.Selection.Find
    find.Text = search_text
    find.Execute()

    time.sleep(2)

    p.press('pgdn')
    time.sleep(1)
    p.hotkey('ctrl', 'v')
    
    time.sleep(10)
    # Salvar e fechar o documento
    doc.SaveAs(file_docx_result)
    doc.Close()
    
    time.sleep(1)
    # Fechar o Word
    word.Quit()

    print('Conteúdo colado com sucesso no novo documento!')

def find_window_by_title(title):
    """ Procura uma janela pelo título e retorna se ela for encontrada """
    windows = gw.getWindowsWithTitle(title)
    return windows[0] if windows else None

def format_entire_document(file_rtf_in, file_docx_out, start_text, end_text):   
    wdCollapseStart = 1
    time.sleep(3)
    # Abrir o Word
    word = win32.Dispatch('Word.Application')
    word.Visible = True
    
    # Abrir o documento especificado
    doc = word.Documents.Open(file_rtf_in)
    cont = 0
    janela = find_window_by_title("PGR")
    # while janela is None:
    #     janela = ahk.find_window_by_title("PGR", title_match_mode=2)
    #     cont +=1
    #     time.sleep(1)
    #     if cont == 60: raise Exception("Timout esperando a janela do arquivo de entrada RTF")
    
    while janela is None:
        print("Janela não encontrada, tentando novamente...")
        time.sleep(1)
        janela = find_window_by_title("PGR")

    time.sleep(3)

    # Acessar a seleção do Word
    selection = doc.Application.Selection

    time.sleep(1)
    # Localizar o texto de início "Setor"
    find = selection.Find
    find.Text = start_text
    find.Execute()
    start_position = selection.Start

    time.sleep(1)
    # Localizar o texto de fim "Matriz de Risco"
    selection.Collapse(Direction=wdCollapseStart)  # Mover o cursor para o início da seleção atual
    find.Text = end_text
    find.Execute()
    end_position = selection.Start

    time.sleep(1)
    # Selecionar o texto entre "Setor" e "Matriz de Risco"
    if start_position < end_position:
        doc.Range(start_position, end_position).Select()
    else:
        # Caso o texto final seja encontrado antes do texto inicial, selecionar até o final do documento
        doc.Range(start_position, doc.Content.End).Select()

    # Alterar a fonte para 'Verdana-Bold'
    time.sleep(2)
    selection = word.Selection
    selection.Font.Name = 'Verdana-Bold'
    time.sleep(1)

    # Alterar o tamanho da fonte para 8
    time.sleep(2)
    selection.Font.Size = 8
    time.sleep(1)

   
    time.sleep(3)
    # Copiar a seleção
    selection.Copy()

    time.sleep(2)
    paste_content_to_new_document(file_docx_out, "INVENTÁRIO DE RISCOS.")

    # Salvar e fechar o documento
    # doc.SaveAs(file_rtf_in)
    time.sleep(1)
    # doc.Close(SaveChanges=False)

    # Fechar o Word
    word.Quit()
    kill_process_word()
    time.sleep(1)
    print('Documento formatado e salvo com sucesso!')


def remove_rows_with_text(doc, text_to_find):
    '''Remove linhas de docx, passando doc, de python-docx'''
    for table in doc.tables:
        rows_to_delete = []
        for i, row in enumerate(table.rows):
            for cell in row.cells:
                if text_to_find in cell.text:
                    rows_to_delete.append(i)
                    break
        for i in reversed(rows_to_delete):
            tbl = table._tbl
            tbl.remove(tbl.tr_lst[i])


def highlight_cells_with_text(doc, text_to_find):
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                if text_to_find in cell.text:
                    # Pinta a célula encontrada
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:val'), 'clear')
                    shading_elm.set(qn('w:color'), 'auto')
                    shading_elm.set(qn('w:fill'), 'A9A9A9')  # Cinza mais escuro
                    cell._element.get_or_add_tcPr().append(shading_elm)

                    # Pinta as células nas duas linhas seguintes
                    for i in range(1, 3):
                        next_row = table.rows[row_idx + i]
                        next_cell = next_row.cells[cell_idx]
                        shading_elm = OxmlElement('w:shd')
                        shading_elm.set(qn('w:val'), 'clear')
                        shading_elm.set(qn('w:color'), 'auto')
                        shading_elm.set(qn('w:fill'), 'A9A9A9')  # Cinza mais escuro
                        next_cell._element.get_or_add_tcPr().append(shading_elm)


def highlight_cells_with_text(doc, text_to_find):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if text_to_find in cell.text:
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:val'), 'clear')
                    shading_elm.set(qn('w:color'), 'auto')
                    shading_elm.set(qn('w:fill'), 'D3D3D3')  # Cor cinza
                    cell._element.get_or_add_tcPr().append(shading_elm)


def find_and_update_table(doc, search_text:str, new_value:str, index:int, dt_now:str):
    for table in doc.tables:
        for row in table.rows:
            first_cell = row.cells[0]
            if search_text in first_cell.text:
                row.cells[index].text = new_value
                row.cells[13].text = dt_now
                

class Receita():
    def __init__(self) -> None:
        pass

    def consulta_cnpj_receita_federal(self, cnpj:str, token:str):
        url = f'https://api.cnpja.com/rfb?taxId={cnpj}'
        res = requests.get(url, headers={"Authorization": token})
        if res.status_code == 200:
            print('consuta na receita ok')
            return res.json()
        elif res.status_code == 429:
            # sem creditos, ou excedeu limite por minuto
            dados = res.json()
            raise Exception(str(dados['message']))
        else:
            msg = f"Erro api da receita status_code: {res.status_code}"
            raise Exception(msg)
        
    def extrair_dados_receita(self, json_receita:dict):
        dt = json_receita
        try: cod_desc_secund = str(dt['sideActivities'][0]['id']) +', '+ str(dt['sideActivities'][0]['text'])
        except: cod_desc_secund = "*****"

        try: mail = str(dt['emails'][0]['address'])
        except: mail = "*****"

        try: tel = str(dt['phones'][0]['area']) +' '+ str(dt['phones'][0]['number'])
        except: tel = "*****"

        try: nome_fantasia = str(dt['alias'])
        except: nome_fantasia = "*****"

        keys_receita = {
            "rec_num_incricao": str(dt['taxId']),
            "rec_dt_abertura": str(dt['founded']),
            "rec_nome_empresa": str(dt['name']),
            "rec_nome_fantasia": nome_fantasia,
            "rec_porte": str(dt['size']['acronym']),
            "rec_cod_desc_mun": str(dt['mainActivity']['id']) +', '+ str(dt['mainActivity']['text']),
            "rec_cod_desc_secund": cod_desc_secund,
            "rec_cod_desc_jur": str(dt['nature']['id']) +', '+ str(dt['nature']['text']),
            "rec_logradouro": str(dt['address']['street']),
            "rec_num": str(dt['address']['number']),
            "rec_complemento": str(dt['address']['details']),
            "rec_cep": str(dt['address']['zip'])[:-3] + '-' + str(dt['address']['zip'])[-3:],
            "rec_bairro": str(dt['address']['district']),
            "rec_municipio": str(dt['address']['city']),
            "rec_uf": str(dt['address']['state']),
            "rec_mail": mail,
            "rec_tel": tel,
            "rec_ente_resp": "*****",
            "rec_situacao_cad": str(dt['status']['text']),
            "rec_dta_situacao": dt['statusDate'],
            "rec_situacao_especial": "*****",
            "rec_dt_situacao_esp": "*****"
            }
        return keys_receita


class Doc_Rtf():

    def doc_rtf_to_dataframe(self, file_base_rtf):
        # Converte o arquivo RTF para HTML
        output = pypandoc.convert_file(file_base_rtf, 'html')
        tables = pd.read_html(StringIO(output), flavor="html5lib")
        return tables
    

    def calc_data_vigencia(self, file_base_rtf:str):
        try:
            file_data = file_base_rtf.split('-')
            ano_ini = str(file_data[11]).strip()
            ano_fim = str(file_data[2]).strip()
            mes = str(file_data[3]).strip().upper()
        except:
            ano_ini = str(datetime.now().year)
            ano_fim = str( (datetime.now().year) + 2)
            mes = datetime.now().strftime('%B').upper()
        data_vigencia = f"{mes} {ano_ini} A {mes} {ano_fim}"
        return data_vigencia
    

    def ler_tabelas_metodo1(self, tables):
        prog_auditivo = False
        prog_respiratorio = False
        prog_NR6 = False
        prog_NR10 = False
        prog_NR11 = False
        prog_NR12 = False
        prog_NR13 = False
        prog_NR33 = False
        prog_NR35 = False
        print("Processando regras de negócio...")

        for tbl in tables:
            if str(tbl[0][0]).strip() == "Identificação": 
                tbl_empresa = tbl
                nome_empresa = str(tbl_empresa[0][1]).replace('Empresa ','').strip()
                cnpj = str(tbl_empresa[2][2]).replace('CNPJ ','').strip()
                cnae = str(tbl_empresa[0][4]).replace('CNAE ','').strip()
                descr_cnae = str(tbl_empresa[2][4]).replace('Descrição CNAE ','').strip()
                grau_risco = str(tbl_empresa[1][4]).replace('Grau de Risco ','').strip()

            elif str(tbl[0][0]).strip() == "Total de Funcionários": 
                tbl_funcionarios = tbl
                masculino = str(tbl_funcionarios[1][2]).strip()
                feminino = str(tbl_funcionarios[2][2]).strip()
                menor = str(tbl_funcionarios[3][2]).strip()
                total = str(tbl_funcionarios[4][2]).strip()
                
            elif str(tbl[0][0]).strip() == 'Agente': 
                agente = str(tbl[1][0]).lower()
                grupo = str(tbl[3][0]).lower().strip()
                
                descricao_epis = "Não aplicado"
                for x in range(len(tbl)):
                    if str(tbl[0][x]).strip() == "Orientação":
                        descricao_epis = str(tbl[1][x]).strip()
                        break
                    elif "EPI's" in str(tbl[1][x]):
                        descricao_epis = str(tbl[1][x]).strip()
                        break

                # testa se é necessário o programa auditivo
                if ("RUÍDO CONTINUO (ACIMA DE 85 dB)".lower() in agente) and (prog_auditivo == False): 
                    prog_auditivo = True   
                elif ("RUÍDO CONTINUO (RUÍDOS ENTRE 80 A 85 dB)".lower() in agente) and (prog_auditivo == False): 
                    prog_auditivo = True
                    
                # testa se é necessário o programa respiratorio
                risco_respiratorio = ["químico", "biológico"]
                if (grupo in risco_respiratorio) and (prog_respiratorio == False): 
                    prog_respiratorio = True

                # testa se é necessário o programa NR6
                risco_nr6 = ["químico", "físico", "acidente", "biológico"]
                if (grupo in risco_nr6) and (prog_NR6 == False): 
                    lista_dispensa_de_epi = ["Não há necessidade", "Não aplicado", "Demais EPI's estão dispensados"]
                    for dispensa_de_epi in lista_dispensa_de_epi:
                        if dispensa_de_epi not in descricao_epis:
                            prog_NR6 = True

                # testa se é necessário o programa NR10
                if (grupo == "acidente") and (prog_NR10 == False):
                    if ("arco elétrico" in agente):
                        prog_NR10 = True

                # testa se é necessário o programa NR11
                if (grupo == "acidente") and (prog_NR11 == False):
                    if ("empilhadeira" in agente) or ("transpaleteira" in agente):
                        prog_NR11 = True

                # testa se é necessário o programa NR12
                if (grupo == "acidente") and (prog_NR12 == False):
                    if ("operação de máquinas e equipamentos" in agente):
                        prog_NR12 = True

                # testa se é necessário o programa NR13
                if (grupo == "acidente") and (prog_NR13 == False):
                    if ("vaso de pressão" in agente):
                        prog_NR13 = True

                # testa se é necessário o programa NR33
                if (grupo == "acidente") and (prog_NR33 == False):
                    if ("espaço confinado" in agente):
                        prog_NR33 = True


                # testa se é necessário o programa NR35
                if (grupo == "acidente") and (prog_NR35 == False):
                    if ("trabalho em altura" in agente):
                        prog_NR35 = True

        dados_result = {
            "prog_auditivo": prog_auditivo,
            "prog_respiratorio": prog_respiratorio,
            "prog_NR6": prog_NR6,
            "prog_NR10": prog_NR10,
            "prog_NR11": prog_NR11,
            "prog_NR12": prog_NR11,
            "prog_NR13": prog_NR13,
            "prog_NR33": prog_NR33,
            "prog_NR35": prog_NR35,
            "nome_empresa": nome_empresa,
            "cnpj": cnpj,
            "cnae": cnae,
            "descr_cnae": descr_cnae,
            "grau_risco": grau_risco,
            "masculino": masculino,
            "feminino": feminino,
            "menor": menor,
            "total": total
            }
        
        return dados_result
    

    def ler_tabelas_metodo2(self, tables):
        prog_auditivo = False
        prog_respiratorio = False
        prog_NR6 = False
        prog_NR10 = False
        prog_NR11 = False
        prog_NR12 = False
        prog_NR13 = False
        prog_NR33 = False
        prog_NR35 = False
        print("Processando regras de negócio...")

        for n in range(len(tables)):
            # tables[num_tabela][coluna][linha]

            if str(tables[n][0][0]).strip() == "Identificação":
                nome_empresa = str(tables[n+1][0][0]).replace('Empresa ','').strip()
                cnpj = str(tables[n+2][2][0]).replace('CNPJ ','').strip()
                cnae = str(tables[n+4][0][0]).replace('CNAE ','').strip()
                descr_cnae = str(tables[n+4][2][0]).replace('Descrição CNAE ','').strip()
                grau_risco = str(tables[n+4][1][0]).replace('Grau de Risco ','').strip()

            elif str(tables[n][0][0]).strip() == "Total de Funcionários": 
                masculino = str(tables[n+3][0][0]).strip()
                feminino = str(tables[n+3][1][0]).strip()
                menor = str(tables[n+3][2][0]).strip()
                total = str(tables[n+3][3][0]).strip()
                
            elif str(tables[n][0][0]).strip() == 'Agente': 
                agente = str(tables[n][1][0]).lower()
                grupo = str(tables[n][3][0]).lower().strip()
                idx = n

                while True:
                    idx +=1
                    if str(tables[idx][0][0]).strip() == "Orientação":
                        descricao_epis = str(tables[idx+1][0][0]).strip()
                        break
                    elif "EPI's" in str(tables[idx][0][0]).strip():
                        descricao_epis = str(tables[idx][0][0]).strip()
                        break
                    elif str(tables[idx][0][0]).strip() == 'Agente':
                        descricao_epis = "Não aplicado"
                        break
                    elif str(tables[idx][0][0]).strip() == "Total de Funcionários":
                        descricao_epis = "Não aplicado"
                        break

                # testa se é necessário o programa auditivo
                if ("RUÍDO CONTINUO (ACIMA DE 85 dB)".lower() in agente) and (prog_auditivo == False): 
                    prog_auditivo = True   
                elif ("RUÍDO CONTINUO (RUÍDOS ENTRE 80 A 85 dB)".lower() in agente) and (prog_auditivo == False): 
                    prog_auditivo = True
                    
                # testa se é necessário o programa respiratorio
                risco_respiratorio = ["químico", "biológico"]
                if (grupo in risco_respiratorio) and (prog_respiratorio == False): 
                    prog_respiratorio = True

                # testa se é necessário o programa NR6
                risco_nr6 = ["químico", "físico", "acidente", "biológico"]
                if (grupo in risco_nr6) and (prog_NR6 == False): 
                    lista_dispensa_de_epi = ["Não há necessidade", "Não aplicado", "Demais EPI's estão dispensados"]
                    for dispensa_de_epi in lista_dispensa_de_epi:
                        if dispensa_de_epi not in descricao_epis:
                            prog_NR6 = True

                # testa se é necessário o programa NR10
                if (grupo == "acidente") and (prog_NR10 == False):
                    if ("arco elétrico" in agente):
                        prog_NR10 = True

                # testa se é necessário o programa NR11
                if (grupo == "acidente") and (prog_NR11 == False):
                    if ("empilhadeira" in agente) or ("transpaleteira" in agente):
                        prog_NR11 = True

                # testa se é necessário o programa NR12
                if (grupo == "acidente") and (prog_NR12 == False):
                    if ("operação de máquinas e equipamentos" in agente):
                        prog_NR12 = True

                # testa se é necessário o programa NR13
                if (grupo == "acidente") and (prog_NR13 == False):
                    if ("vaso de pressão" in agente):
                        prog_NR13 = True

                # testa se é necessário o programa NR33
                if (grupo == "acidente") and (prog_NR33 == False):
                    if ("espaço confinado" in agente):
                        prog_NR33 = True


                # testa se é necessário o programa NR35
                if (grupo == "acidente") and (prog_NR35 == False):
                    if ("trabalho em altura" in agente):
                        prog_NR35 = True

        dados_result = {
            "prog_auditivo": prog_auditivo,
            "prog_respiratorio": prog_respiratorio,
            "prog_NR6": prog_NR6,
            "prog_NR10": prog_NR10,
            "prog_NR11": prog_NR11,
            "prog_NR12": prog_NR11,
            "prog_NR13": prog_NR13,
            "prog_NR33": prog_NR33,
            "prog_NR35": prog_NR35,
            "nome_empresa": nome_empresa,
            "cnpj": cnpj,
            "cnae": cnae,
            "descr_cnae": descr_cnae,
            "grau_risco": grau_risco,
            "masculino": masculino,
            "feminino": feminino,
            "menor": menor,
            "total": total
            }
        
        return dados_result


    def __init__(self, file_base_rtf:str):
        kill_process_word()
        time.sleep(1)
        self.file_base_rtf = file_base_rtf
        self.tables = self.doc_rtf_to_dataframe(file_base_rtf)
        self.data_vigencia = self.calc_data_vigencia(file_base_rtf)
        self.data_elaboracao = datetime.now().strftime('%d.%m.%Y')
        self.data_conclusao = datetime.now().strftime('%d %B DE %Y').upper()
        # programas de conservacao
        try:
            self.dados_result = self.ler_tabelas_metodo1(self.tables)
        except:
            try:
                self.dados_result = self.ler_tabelas_metodo2(self.tables)
            except:
                raise Exception("Erro, Nao foi possivel organizar os dados do documento de entrada, possivelmente a estrutura ficou diferente ao transformar em HTML")

        # Dados da Empresa
        self.nome_empresa = str(self.dados_result['nome_empresa']).replace('Empresa ','').strip()
        self.cnpj = str(self.dados_result['cnpj']).replace('CNPJ ','').strip()
        self.cnae = str(self.dados_result['cnae']).replace('CNAE ','').strip()
        self.descr_cnae = str(self.dados_result['descr_cnae']).replace('Descrição CNAE ','').strip()

        # Grau de Risco
        self.grau_risco = str(self.dados_result['grau_risco']).replace('Grau de Risco ','').strip()
        if self.grau_risco == '1':
            self.gr1, self.gr2, self.gr3, self.gr4 = '1', 'g.risc', 'g.risc', 'g.risc'
        if self.grau_risco == '2':
            self.gr1, self.gr2, self.gr3, self.gr4 = 'g.risc', '2', 'g.risc', 'g.risc'
        if self.grau_risco == '3':
            self.gr1, self.gr2, self.gr3, self.gr4 = 'g.risc', 'g.risc', '3', 'g.risc'
        if self.grau_risco == '4':
            self.gr1, self.gr2, self.gr3, self.gr4 = 'g.risc', 'g.risc', 'g.risc', '4'


        # Qtd Funcionarios
        self.masculino = str(self.dados_result['masculino']).strip()
        self.feminino = str(self.dados_result['feminino']).strip()
        self.menor = str(self.dados_result['menor']).strip()
        self.total = str(self.dados_result['total']).strip()

        # Referencia para pintar celular qtd funcionarios CIPA
        tt = int(self.total)
        if tt <= 19: self.ref_cipa = "0 a 19"
        elif tt >  19  and tt <  30:  self.ref_cipa = "20 a 29"
        elif tt >  29  and tt <  51:  self.ref_cipa = "30 a 50"
        elif tt >  50  and tt <  81:  self.ref_cipa = "51 a 80"
        elif tt >  80  and tt < 101:  self.ref_cipa = "81 a 100"
        elif tt > 100  and tt < 121:  self.ref_cipa = "101 a 120"
        elif tt > 120  and tt < 141:  self.ref_cipa = "121 a 140"
        elif tt > 140  and tt < 301:  self.ref_cipa = "141 a 300"
        elif tt > 300  and tt < 501:  self.ref_cipa = "301 a 500"
        elif tt > 500  and tt < 1001: self.ref_cipa = "501 a 1000"
        elif tt > 1000 and tt < 2501: self.ref_cipa = "1001 a 2500"
        elif tt > 2500 and tt < 5001: self.ref_cipa = "2501 a 5000"
        elif tt > 5000 and tt <= 10000: self.ref_cipa = "5001 a 10.000"
        elif tt > 10000: self.ref_cipa = "Acima de 10.000"

        # Programas de conservação
        self.prog_auditivo = bool(self.dados_result['prog_auditivo'])
        self.prog_respiratorio = bool(self.dados_result['prog_respiratorio'])
        self.prog_NR6 = bool(self.dados_result['prog_NR6'])
        self.prog_NR10 = bool(self.dados_result['prog_NR10'])
        self.prog_NR11 = bool(self.dados_result['prog_NR11'])
        self.prog_NR12 = bool(self.dados_result['prog_NR12'])
        self.prog_NR13 = bool(self.dados_result['prog_NR13'])
        self.prog_NR33 = bool(self.dados_result['prog_NR33'])
        self.prog_NR35 = bool(self.dados_result['prog_NR35'])

        # Numero de Integrantes necessários CIPA


        self.keys_rtf = {
            "ref_nome_empresa": self.nome_empresa,
            "ref_data_vigencia": self.data_vigencia,
            "ref_campo_data": self.data_elaboracao,
            "ref_data_conclusao": self.data_conclusao,
            "ref_cnae": self.cnae,
            "ref_grau_risco": self.grau_risco,
            "ref_descr_cnae": self.descr_cnae,
            "ref_masculino": self.masculino,
            "ref_feminino": self.feminino,
            "ref_menor": self.menor,
            "ref_total": self.total,
            "g.risc1": self.gr1,
            "g.risc2": self.gr2,
            "g.risc3": self.gr3,
            "g.risc4": self.gr4
        }


    
