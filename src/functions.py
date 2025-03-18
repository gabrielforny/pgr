
from datetime import datetime
# BIBLIOTECAS
import pandas as pd
import pypandoc
pypandoc.download_pandoc()

from io import StringIO
import locale, time
import pythoncom
import os
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import locale
locale.setlocale(locale.LC_ALL, 'pt_BR.utf8')
import psutil
import subprocess
import requests

import pygetwindow as gw

import pyautogui as p

import fitz  # PyMuPDF
from PIL import Image
from docx.shared import Inches
from docx import Document
import re
import pypandoc

import os
import pythoncom
import traceback

import win32com.client as win32
import pywintypes
import pyautogui
import win32clipboard

def kill_process_word():
    process_name = "WINWORD"
    # Verificar se o processo está em execução
    for proc in psutil.process_iter(['name']):
        if process_name.lower() in proc.info['name'].lower():
            # Encerrar o processo
            pid = proc.pid
            subprocess.run(['taskkill', '/PID', str(pid), '/F'])

def find_window_by_title(title):
    """ Procura uma janela pelo título e retorna se ela for encontrada """
    windows = gw.getWindowsWithTitle(title)
    return windows[0] if windows else None

def paste_content_to_new_document(file_docx_result, search_text):
    # Abrir o Word
    time.sleep(3)
    word = win32.Dispatch('Word.Application')
    word.Visible = True
    
    # Abrir o documento alvo
    doc = word.Documents.Open(file_docx_result)
    time.sleep(1)
    
    # Esperar pela janela com o título parcial "result"
    cont = 0
    # janela = ahk.find_window_by_title("PGR", title_match_mode=2)
    janela = find_window_by_title("PGR")
    
    # while janela is None:
    #     janela = ahk.find_window_by_title("PGR", title_match_mode=2)
    #     cont +=1
    #     time.sleep(1)
    #     if cont == 60: raise Exception("Timout esperando a janela do arquivo result PGR")
    
    while janela is None:
        print("Janela não encontrada, tentando novamente...")
        time.sleep(1)
        janela = find_window_by_title("PGR")

    # Localizar o texto "INVENTÁRIO DE RISCOS."
    time.sleep(1)
    find = word.Selection.Find
    find.Text = search_text
    find.Execute()

    time.sleep(2)

    p.press('pgdn')
    time.sleep(1)
    p.hotkey('ctrl', 'v')
    
    time.sleep(10)
    # Salvar e fechar o documento
    doc.SaveAs(file_docx_result)
    doc.Close()
    
    time.sleep(1)
    # Fechar o Word
    word.Quit()

    print('Conteúdo colado com sucesso no novo documento!')

def formatar_e_inserir_conteudo_direto(file_base_rtf, pgr_destino):
    pythoncom.CoInitialize()  # Inicializa o COM apenas uma vez para o processo inteiro

    word = None
    try:
        word = win32.Dispatch('Word.Application')
        word.Visible = False
        word.DisplayAlerts = 0  # Desativa completamente todas as mensagens do Word

        # Abre o documento base (RTF) e executa a exclusão e formatação de tabelas
        doc = word.Documents.Open(file_base_rtf)
        tables_count = doc.Tables.Count

        # Exclui as primeiras cinco tabelas e a última, se houver mais de uma tabela
        if tables_count >= 5:
            for _ in range(5):
                doc.Tables(1).Delete()
        if doc.Tables.Count > 0:
            doc.Tables(doc.Tables.Count).Delete()

        # Remover tabelas com linhas em branco (ignorando erros de mesclagem)
        tables_to_remove = []
        for i in range(1, doc.Tables.Count + 1):
            try:
                table = doc.Tables(i)
                blank_row_found = False
                for row in table.Rows:
                    row_text = ''.join([cell.Range.Text.strip() for cell in row.Cells])
                    if not row_text:
                        blank_row_found = True
                        break
                if blank_row_found:
                    tables_to_remove.append(i)
            except pywintypes.com_error:
                print(f"Ignorando tabela {i} devido a células mescladas verticalmente.")
                continue
        for index in reversed(tables_to_remove):
            doc.Tables(index).Delete()

        # Formatar todas as células das tabelas restantes
        for table in doc.Tables:
            try:
                for row in table.Rows:
                    for cell in row.Cells:
                        cell_range = cell.Range
                        cell_range.Font.Name = 'Verdana'
                        cell_range.Font.Size = 8
            except pywintypes.com_error:
                print("Erro ao formatar células devido a células mescladas; continuando com outras tabelas.")
                continue

        # Reorganizar o documento para inserção de quebras de página
        range_total = doc.Range()
        paragraphs = range_total.Paragraphs
        first_setor = True
        cargo_encontrado = False
        setor_atual = ""

        for paragraph in paragraphs:
            texto = paragraph.Range.Text.strip()
            if texto.startswith("Setor:"):
                if not first_setor and texto != setor_atual:
                    paragraph.Range.InsertBreak(7)
                setor_atual = texto
                first_setor = False
                cargo_encontrado = False
                continue
            if texto.startswith("Cargo:"):
                if not cargo_encontrado:
                    cargo_encontrado = True
                else:
                    paragraph.Range.InsertBreak(7)

        # Seleciona e copia o conteúdo reorganizado
        range_total.Select()
        word.Selection.Copy()
        doc.Close(SaveChanges=False)  # Fecha o documento RTF sem salvar

        # Abre o documento de destino (DOCX) e insere o conteúdo copiado no marcador especificado
        destino_doc = word.Documents.Open(pgr_destino)
        destino_range = destino_doc.Content
        placeholder = "{{tabela}}"
        find_placeholder = destino_range.Find
        find_placeholder.Text = placeholder

        # Tenta localizar o marcador e colar o conteúdo
        if find_placeholder.Execute():
            print(f"Marcador {placeholder} encontrado no documento de destino. Colando a tabela.")
            find_placeholder.Parent.Paste()  # Cola o conteúdo copiado
        else:
            print(f"Marcador {placeholder} não encontrado no documento {pgr_destino}.")

        # Remover as duas últimas tabelas
        table_count = destino_doc.Tables.Count
        if table_count >= 2:
            destino_doc.Tables(table_count).Delete()
            destino_doc.Tables(table_count - 1).Delete()
            print("As duas últimas tabelas foram removidas.")

        # Salva o documento final com o mesmo nome de `pgr_destino`
        destino_doc.SaveAs(pgr_destino)
        destino_doc.Close()
        print(f"Documento final salvo em: {pgr_destino}")

    except pywintypes.com_error as e:
        print(f"Erro ao processar o documento: {e}")
        traceback.print_exc()
    finally:
        pythoncom.CoUninitialize()
        
        if word:
            word.Quit()

from docx import Document

def remove_paginas_vazias_rapido(doc_path, output_path):
    doc = Document(doc_path)
    novo_doc = Document()

    # Lista de seções para verificar páginas separadamente
    secoes = doc.sections

    # Armazena o conteúdo de cada seção para verificar se está vazia
    conteudo_geral = []
    
    for paragrafo in doc.paragraphs:
        texto = paragrafo.text.strip()
        
        # Se o texto não for vazio, adiciona à lista
        if texto:
            conteudo_geral.append(paragrafo.text)

    # Se o documento inteiro está vazio, salva um documento em branco e retorna
    if not conteudo_geral:
        novo_doc.add_paragraph("Documento sem conteúdo válido")
        novo_doc.save(output_path)
        print("O documento estava completamente vazio. Arquivo salvo sem páginas em branco.")
        return

    # Separamos as seções e checamos se há conteúdo em cada uma
    for secao in secoes:
        if any(par.text.strip() for par in doc.paragraphs):
            for parag in doc.paragraphs:
                novo_doc.add_paragraph(parag.text)

    # Salvar o novo documento
    novo_doc.save(output_path)
    print(f"Documento atualizado salvo em: {output_path}")

def atualizar_indice(doc_path):
    pythoncom.CoInitialize()
    word = win32.Dispatch("Word.Application")
    word.Visible = False
    doc = None

    try:
        doc = word.Documents.Open(doc_path)
        if doc.TablesOfContents.Count == 0:
            print(f"Nenhuma Tabela de Conteúdos encontrada em {doc_path}.")
            return

        for toc in doc.TablesOfContents:
            toc.Update()

            for paragraph in toc.Range.Paragraphs:
                for run in paragraph.Range.Words:
                    if run.Text.strip(): 
                        run.Font.Name = "Verdana"
                        run.Font.Size = 8
                        run.Font.Bold = True
                        run.Font.ColorIndex = 1 
                        paragraph.Format.SpaceBefore = 0
                        paragraph.Format.SpaceAfter = 0
                        paragraph.Format.LineSpacingRule = 0

            toc.UpdatePageNumbers()

        doc.Fields.Update()
        doc.Save()
        doc.Close()
        print(f"Tabela de Conteúdos formatada e documento salvo com sucesso em {doc_path}.")

    except Exception as e:
        print(f"Ocorreu um erro no arquivo {doc_path}: {e}")
    finally:
        pythoncom.CoUninitialize()

def exportar_para_pdf(doc_path, pdf_destino):
    try:
        pythoncom.CoInitialize()
        word = win32.Dispatch('Word.Application')
        word.Visible = False  # Mantém o Word em segundo plano
        doc = word.Documents.Open(doc_path)  # Abre o documento

        # Exporta o documento para o formato PDF
        doc.ExportAsFixedFormat(pdf_destino, 17)  # 17 é o código para PDF no Word
        print(f"Documento exportado para PDF com sucesso. Local: {pdf_destino}")

        doc.Close(False)  # Fecha o documento sem salvar
        time.sleep(3)
        word.Quit()  # Fecha o Word

    except Exception as e:
        print(f"Erro ao exportar para PDF: {e}")
    finally:
        pythoncom.CoUninitialize()

def copiar_plano_de_acao(file_base_rtf, pgr_destino):
    pythoncom.CoInitialize()

    word = None
    try:
        # Inicializar o Word e abrir o arquivo RTF
        word = win32.Dispatch('Word.Application')
        word.Visible = False  # Deixa o Word visível
        word.DisplayAlerts = 0

        # Abrir o arquivo de origem
        doc_rtf = word.Documents.Open(file_base_rtf)

        # Procurar o termo "Todas"
        find = doc_rtf.Content.Find
        find.Text = "Todas"
        find.Execute()

        if not find.Found:
            raise ValueError("'Todas' não foi encontrado no arquivo base.")

        print("Termo 'Todas' encontrado. Selecionando conteúdo...")

        # Definir o range para pegar a linha anterior ao termo 'Todas' até o fim do documento
        start_position = find.Parent.Paragraphs(1).Range.Start - 1  # Começa uma linha antes
        end_position = doc_rtf.Content.End
        content_range = doc_rtf.Range(Start=start_position, End=end_position)

        # Copiar o conteúdo selecionado
        content_range.Copy()
        print("Conteúdo copiado com sucesso.")

        # Fechar o arquivo RTF sem salvar
        doc_rtf.Close(False)

        # Abrir o arquivo de destino
        doc_destino = word.Documents.Open(pgr_destino)
        word.Visible = False

        # Procurar o marcador "INFO_PLANO DE AÇÃO"
        find_destino = doc_destino.Content.Find
        find_destino.Text = "INFO_PLANO DE AÇÃO"
        find_destino.Execute()

        if find_destino.Found:
            print("'INFO_PLANO DE AÇÃO' encontrado no arquivo de destino. Substituindo conteúdo.")
            destino_range = find_destino.Parent
            destino_range.Text = "" 
            time.sleep(3)

            try:
                print("Tentando colar com PasteSpecial...")
                destino_range.Paste()
            except pywintypes.com_error:
                print("PasteSpecial falhou, tentando Paste normal...")
                destino_range.PasteSpecial(DataType=22)
                
            time.sleep(1) 
            print("Conteúdo colado com sucesso. - PLANO DE AÇÃO")

            # Quebrar em novas páginas sempre que encontrar "NR" ou "Treinamento"
            paragraphs = doc_destino.Paragraphs
            found_atividade = False  # Variável para controlar se já encontrou o "Atividade"
            skip_first_break = False  # Variável para controlar se deve pular a quebra após "Atividade"

            for para in paragraphs:
                text = para.Range.Text.strip()
                
                # Verificar se encontramos "Atividade"
                if text.startswith("Atividade"):
                    found_atividade = True  # Marca que encontramos "Atividade"
                    continue  # Pula para o próximo parágrafo
                
                # Verificar se encontramos "NR" ou "Treinamento"
                if (text.startswith("NR") or text.startswith("Treinamento")) and text != "Treinamento e competência":
                    # Se for o primeiro "NR" ou "Treinamento" após "Atividade", não quebra a página
                    if found_atividade and not skip_first_break:
                        skip_first_break = True  # Marca que o próximo "NR" ou "Treinamento" não vai quebrar a página
                    else:
                        para.Range.InsertBreak(7)  # 7 representa wdParagraphBreakPageBreak

        else:
            raise ValueError("'INFO_PLANO DE AÇÃO' não foi encontrado no documento de destino.")

        # Salvar e fechar o documento de destino
        doc_destino.SaveAs(pgr_destino)
        time.sleep(3)
        doc_destino.Close()

        print("Conteúdo copiado e colado com sucesso! - PLANO DE AÇÃO")

    except Exception as e:
        print(f"Ocorreu um erro: {e}")

    finally:
        pythoncom.CoUninitialize()
        return pgr_destino
        
        # if word:
        #     word.Quit()

def copiar_inventario_via_range(file_base_rtf, pgr_destino, nome_empresa):
    pythoncom.CoInitialize()

    word = None
    try:
        # Inicializar o Word e abrir o arquivo RTF
        word = win32.Dispatch('Word.Application')
        word.Visible = False  # Deixa o Word visível
        word.DisplayAlerts = 0

        # Abrir o arquivo de origem
        doc_rtf = word.Documents.Open(file_base_rtf)

        unidade_account = 0
        find = doc_rtf.Content.Find
        find.Text = "Caracterização"
        find.MatchCase = False
        ranges = []
        
        while find.Execute(): 
            unidade_account += 1
            frist_range = find.Parent
            ranges.append(find.Parent) 

        if unidade_account < 1:
            raise ValueError("Menos de 1 ocorrências de 'Caracterização dos processos e ambientes de trabalho' encontradas no arquivo base.")

        print(f"{unidade_account} ocorrências de '{nome_empresa}' encontradas....")

        # Usar o 8º resultado
        start_position = frist_range.Start 

        plano_de_count = 0
        findPlano = doc_rtf.Content.Find
        findPlano.Text = "Plano de"  # Procurar "Plano de"
        findPlano.MatchCase = False  # Ignorar maiúsculas/minúsculas
        while findPlano.Execute():  # Executar até o final
            plano_de_count += 1
            last_plano_range = findPlano.Parent  # Armazenar a última ocorrência de "Plano de"

        if plano_de_count == 0:
            raise ValueError("'Plano de' não foi encontrado no arquivo base.")

        print(f"{plano_de_count} ocorrências de 'Plano de' encontradas. Usando a última...")

        # Agora, usar a última posição encontrada de "Plano de"
        end_position = last_plano_range.Start
        # Definir o range de conteúdo a ser copiado
        content_range = doc_rtf.Range(Start=start_position, End=end_position)

        # Copiar o conteúdo selecionado
        content_range.Copy()
        print("Conteúdo copiado com sucesso. - Inventário de riscos")

        # Fechar o arquivo RTF
        doc_rtf.Close(False)

        # Abrir o arquivo de destino
        doc_destino = word.Documents.Open(pgr_destino)
        word.Visible = False

        # Procurar o marcador "INFO_INVENTÁRIO DE RISCOS"
        find_destino = doc_destino.Content.Find
        find_destino.Text = "INFO_INVENTÁRIO DE RISCOS"
        find_destino.Execute()

        if find_destino.Found:
            print("'INFO_INVENTÁRIO DE RISCOS' encontrado no arquivo de destino. Substituindo conteúdo.")
            destino_range = find_destino.Parent
            destino_range.Text = ""  # Limpar o conteúdo existente
            time.sleep(5)  # Pausa para garantir que o conteúdo seja colado corretamente

            # Colar o conteúdo copiado mantendo a formatação
            try:
                print("Tentando colar com PasteSpecial...")
                destino_range.Paste()
            except pywintypes.com_error:
                print("PasteSpecial falhou, tentando Paste normal...")
                destino_range.PasteSpecial(DataType=22)

            time.sleep(1)  # Pausa para garantir que o conteúdo seja colado corretamente
            print("Conteúdo copiado e colado com sucesso.")

            # Quebrar em novas páginas com base em "SETOR" e "CARGO"
            paragraphs = doc_destino.Paragraphs
            found_setor = False
            skip_next_cargo = False  # Controle para pular a quebra de página no primeiro "CARGO" após "SETOR"

            for para in paragraphs:
                text = para.Range.Text.strip()

                if text.startswith("SETOR"):
                    # Sempre que encontrar "SETOR", quebra de página
                    para.Range.InsertBreak(7)  # 7 representa wdPageBreak
                    found_setor = True
                    skip_next_cargo = False  # Reinicia o controle de pular o "CARGO"

                elif text.startswith("CARGO"):
                    if found_setor and not skip_next_cargo:
                        # Se o "CARGO" vier logo após o "SETOR", não quebra a página
                        skip_next_cargo = True
                    else:
                        # Caso contrário, quebra a página
                        para.Range.InsertBreak(7)  # 7 representa wdPageBreak
                        found_setor = False  # Reinicia o controle para "SETOR"

        else:
            raise ValueError("'INFO_INVENTÁRIO DE RISCOS' não foi encontrado no documento de destino.")

        # Salvar e fechar o documento de destino
        doc_destino.SaveAs(pgr_destino)
        time.sleep(3)
        doc_destino.Close()

        print("Conteúdo copiado e colado com sucesso!")

    except Exception as e:
        print(f"Ocorreu um erro: {e}")

    finally:
        pythoncom.CoUninitialize()
        return pgr_destino
        
        # if word:
        #     word.Quit()

def remove_rows_with_text(doc, text_to_find):
    '''Remove linhas de docx, passando doc, de python-docx'''
    for table in doc.tables:
        rows_to_delete = []
        for i, row in enumerate(table.rows):
            for cell in row.cells:
                if text_to_find in cell.text:
                    rows_to_delete.append(i)
                    break
        for i in reversed(rows_to_delete):
            tbl = table._tbl
            tbl.remove(tbl.tr_lst[i])


def highlight_cells_with_text(doc, text_to_find):
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                normalized_cell_text = cell.text.replace('\n', ' ').strip()
                
                if text_to_find in normalized_cell_text:
                    # Pinta a célula encontrada
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:val'), 'clear')
                    shading_elm.set(qn('w:color'), 'auto')
                    shading_elm.set(qn('w:fill'), 'A9A9A9')  # Cinza mais escuro
                    cell._element.get_or_add_tcPr().append(shading_elm)

                    # Pinta as células nas duas linhas seguintes, se existirem
                    for i in range(1, 3):
                        if row_idx + i < len(table.rows):  # Verifica se a linha existe
                            next_row = table.rows[row_idx + i]
                            if cell_idx < len(next_row.cells):  # Verifica se a célula existe
                                next_cell = next_row.cells[cell_idx]
                                shading_elm = OxmlElement('w:shd')
                                shading_elm.set(qn('w:val'), 'clear')
                                shading_elm.set(qn('w:color'), 'auto')
                                shading_elm.set(qn('w:fill'), 'A9A9A9')  # Cinza mais escuro
                                next_cell._element.get_or_add_tcPr().append(shading_elm)


def find_and_update_table(doc, search_text:str, new_value:str, index:int, dt_now:str):
    for table in doc.tables:
        for row in table.rows:
            first_cell = row.cells[0]
            if search_text in first_cell.text:
                row.cells[index].text = new_value
                row.cells[13].text = dt_now           

def replace_text_with_images(pdf_path, docx_path, remove_header=True):
    # Configurações
    texts_to_remove = [
        "A avaliação ergonômica preliminar",
        "De acordo com a NR-1, é responsabilidade"
    ]

    # Abrir documentos
    doc_pdf = fitz.open(pdf_path)

    # 1. Localizar e limpar parágrafos alvo
    target_paragraphs = []
    for i, para in enumerate(docx_path.paragraphs):
        for text in texts_to_remove:
            if text in para.text:
                para.clear()  # Apaga o conteúdo do parágrafo
                target_paragraphs.append((i, para))  # Armazena índice e parágrafo

    # Se não encontrar parágrafos para substituir, sair
    if not target_paragraphs:
        print("⚠ Nenhum parágrafo encontrado para substituição!")
        return

    # 2. Processar cada página do PDF
    temp_images = []
    for page_num in range(len(doc_pdf)):
        page = doc_pdf[page_num]
        
        # Renderizar página como imagem com DPI otimizado
        pix = page.get_pixmap(dpi=150)  
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

        # Reduzir consumo de memória convertendo para RGB
        img = img.convert("RGB")

        # Recorte personalizado (ajuste conforme necessidade)
        if remove_header:
            header_height = int(img.height * 0.05)  # 10% para cabeçalho
            footer_height = int(img.height * 0.05)  # 5% para rodapé
            img = img.crop((0, header_height, img.width, img.height - footer_height))

        # Redimensionar mantendo proporção de forma otimizada
        target_width = 500  # Largura em pixels para evitar conversão de Inches
        img.thumbnail((target_width, img.height), Image.LANCZOS)

        # Salvar imagem temporária
        img_path = f"temp_pdf_page_{page_num}.png"
        img.save(img_path, format="PNG", optimize=True)  # Otimizar PNG
        temp_images.append(img_path)

    # 3. Substituir parágrafos limpos pelas imagens
    index_to_insert = target_paragraphs[0][0]  # Índice do primeiro parágrafo removido

    for img_path in temp_images:
        if index_to_insert < len(docx_path.paragraphs):
            new_para = docx_path.paragraphs[index_to_insert]  # Usa o mesmo índice
        else:
            new_para = docx_path.add_paragraph()  # Se não houver mais parágrafos, cria um novo
        
        run = new_para.add_run()
        run.add_picture(img_path, width=Inches(6.5))

        index_to_insert += 1  # Avança para o próximo local de inserção

    # 4. Limpeza final
    for img_path in temp_images:
        if os.path.exists(img_path):
            os.remove(img_path)

    print(f"✅ Documento processado com sucesso")
       
class Receita():
    def __init__(self) -> None:
        pass

    def consulta_cnpj_receita_federal(self, cnpj:str, token:str):
        url = f'https://api.cnpja.com/rfb?taxId={cnpj}'
        res = requests.get(url, headers={"Authorization": token})
        if res.status_code == 200:
            print('consuta na receita ok')
            return res.json()
        elif res.status_code == 429:
            # sem creditos, ou excedeu limite por minuto
            dados = res.json()
            raise Exception(str(dados['message']))
        else:
            msg = f"Erro api da receita status_code: {res.status_code}"
            raise Exception(msg)
        
    def extrair_dados_receita(self, json_receita:dict):
        dt = json_receita
        try: cod_desc_secund = str(dt['sideActivities'][0]['id']) +', '+ str(dt['sideActivities'][0]['text'])
        except: cod_desc_secund = "*****"

        try: mail = str(dt['emails'][0]['address'])
        except: mail = "*****"

        try: tel = str(dt['phones'][0]['area']) +' '+ str(dt['phones'][0]['number'])
        except: tel = "*****"

        try: 
            nome_fantasia = str(dt['alias'])
            
            if nome_fantasia == 'None':
                nome_fantasia = str(dt['name']).replace('.', '')
                
        except: nome_fantasia = "*****"

        keys_receita = {
            "rec_num_incricao": str(dt['taxId']),
            "rec_dt_abertura": str(dt['founded']),
            "rec_nome_empresa": str(dt['name']),
            "rec_nome_fantasia": nome_fantasia,
            "rec_porte": str(dt['size']['acronym']),
            "rec_cod_desc_mun": str(dt['mainActivity']['id']) +', '+ str(dt['mainActivity']['text']),
            "rec_cod_desc_secund": cod_desc_secund,
            "rec_cod_desc_jur": str(dt['nature']['id']) +', '+ str(dt['nature']['text']),
            "rec_logradouro": str(dt['address']['street']),
            "rec_num": str(dt['address']['number']),
            "rec_complemento": str(dt['address']['details']),
            "rec_cep": str(dt['address']['zip'])[:-3] + '-' + str(dt['address']['zip'])[-3:],
            "rec_bairro": str(dt['address']['district']),
            "rec_municipio": str(dt['address']['city']),
            "rec_uf": str(dt['address']['state']),
            "rec_mail": mail,
            "rec_tel": tel,
            "rec_ente_resp": "*****",
            "rec_situacao_cad": str(dt['status']['text']),
            "rec_dta_situacao": dt['statusDate'],
            "rec_situacao_especial": "*****",
            "rec_dt_situacao_esp": "*****"
            }
        return keys_receita


class Doc_Rtf():

    def doc_rtf_to_dataframe(self, file_base_rtf):
        # Converte o arquivo RTF para HTML
        output = pypandoc.convert_file(file_base_rtf, 'html')
        tables = pd.read_html(StringIO(output), flavor="html5lib")
        return tables
    

    def calc_data_vigencia(self, file_base_rtf:str):
        try:
            file_data = file_base_rtf.split('-')
            ano_ini = str(file_data[11]).strip()
            ano_fim = str(file_data[2]).strip()
            mes = str(file_data[3]).strip().upper()
        except:
            ano_ini = str(datetime.now().year)
            ano_fim = str( (datetime.now().year) + 2)
            mes = datetime.now().strftime('%B').upper()
        data_vigencia = f"{mes} {ano_ini} A {mes} {ano_fim}"
        return data_vigencia
    
    def clean_text(self, text):
        """Remove caracteres estranhos, palavras repetidas e espaços extras."""
        text = re.sub(r"[-=—_]", "", text)  # Remove traços e sublinhados
        text = re.sub(r"\s*\r?\n\s*", " ", text)  # Substitui múltiplas quebras de linha por espaço
        text = re.sub(r"[^\w\sÀ-ÿ]", "", text)  # Remove caracteres não alfanuméricos, mantendo acentos
        text = re.sub(r"\s+", " ", text).strip()  # Remove espaços extras
        return text

    def remove_repeticoes(self, text):
        """Remove repetições seguidas no texto."""
        palavras = text.split()
        texto_limpo = []
        
        for palavra in palavras:
            if palavra not in texto_limpo:
                texto_limpo.append(palavra)

        return " ".join(texto_limpo)

    def ler_tabelas_metodo1(self, tables, file_base_rtf):
        prog_auditivo = False
        prog_respiratorio = False
        prog_NR6 = False
        prog_NR10 = False
        prog_NR11 = False
        prog_NR12 = False
        prog_NR13 = False
        prog_NR33 = False
        prog_NR35 = False
        print("Processando regras de negócio...")
    
        
        for tbl in tables:
            if str(tbl[0][0]).strip() == "Identificação":
                    text = pypandoc.convert_file(file_base_rtf, 'plain')

                    # Regex para capturar os dados corretamente
                    bloco_unidade = re.search(r"UNIDADE\s*\n(.*?)(?=\nEndereço|\Z)", text, re.DOTALL)
                    if bloco_unidade:
                        bloco_unidade_text = bloco_unidade.group(1).strip()
                        
                        cnpj_match = re.search(r"(\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2})", bloco_unidade_text)

                        # Pegamos todas as linhas antes do CNPJ (nome da empresa)
                        if cnpj_match:
                            cnpj = cnpj_match.group(1)
                            nome_empresa_com_caracteres = bloco_unidade_text.split(cnpj)[0].strip()
                            nome_empresa = self.remove_repeticoes(self.clean_text(nome_empresa_com_caracteres)).replace('Tí ','')
                        else:
                            cnpj = "Não encontrado"
                            nome_empresa = "Não encontrado"
                            
                        text = self.clean_text(text).replace('Tí ', '')
                        
                        pattern = r'CNAE\s*(\d{6,7})\s*([\w\sÀ-ÿ]+?)(?=\s*Grau de Risco|\Z)'
                        match = re.search(pattern, self.clean_text(text))

                        if match:
                            cnae = match.group(1)
                            cnae_desc = match.group(2).strip() 
                        else:
                            cnae = cnae_desc = ("CNAE não encontrado")
                            
                    # Extraindo Grau de Risco
                    grau_risco_match = re.search(r"Grau de Risco (\d+)", text)

                    dados_extraidos = {
                        "Nome da Empresa": nome_empresa,
                        "CNPJ": cnpj,
                        "CNAE": cnae,
                        "Descrição CNAE": cnae_desc,
                        "Grau de Risco": grau_risco_match.group(1) if grau_risco_match else "Não encontrado"
                    }

                    #Obter os valores de funcionários, caso tenha
                    pattern = r"(\d+)\s*funcionários\s*(\d+)\s*homem\s*(\d+)\s*mulheres\s*(\d+)\s*menores"

                    # Aplica a regex ao texto
                    match = re.search(pattern, text)

                    # Se houver um match, extrai os valores, caso contrário, atribui None
                    if match:
                        total = match.group(1)
                        masculino = match.group(2)
                        feminino = match.group(3)
                        menor = match.group(4)
                    else:
                        total = masculino = feminino = menor = None
                
            elif str(tbl[0][0]).strip() == 'Agente': 
                agente = str(tbl[1][0]).lower()
                grupo = str(tbl[3][0]).lower().strip()
                
                descricao_epis = "Não aplicado"
                for x in range(len(tbl)):
                    if str(tbl[0][x]).strip() == "Orientação":
                        descricao_epis = str(tbl[1][x]).strip()
                        break
                    elif "EPI's" in str(tbl[1][x]):
                        descricao_epis = str(tbl[1][x]).strip()
                        break

                # testa se é necessário o programa auditivo
                if ("RUÍDO CONTINUO (ACIMA DE 85 dB)".lower() in agente) and (prog_auditivo == False): 
                    prog_auditivo = True   
                elif ("RUÍDO CONTINUO (RUÍDOS ENTRE 80 A 85 dB)".lower() in agente) and (prog_auditivo == False): 
                    prog_auditivo = True
                    
                # testa se é necessário o programa respiratorio
                risco_respiratorio = ["químico", "biológico"]
                if (grupo in risco_respiratorio) and (prog_respiratorio == False): 
                    prog_respiratorio = True

                # testa se é necessário o programa NR6
                risco_nr6 = ["químico", "físico", "acidente", "biológico"]
                if (grupo in risco_nr6) and (prog_NR6 == False): 
                    lista_dispensa_de_epi = ["Não há necessidade", "Não aplicado", "Demais EPI's estão dispensados"]
                    for dispensa_de_epi in lista_dispensa_de_epi:
                        if dispensa_de_epi not in descricao_epis:
                            prog_NR6 = True

                # testa se é necessário o programa NR10
                if (grupo == "acidente") and (prog_NR10 == False):
                    if ("arco elétrico" in agente):
                        prog_NR10 = True

                # testa se é necessário o programa NR11
                if (grupo == "acidente") and (prog_NR11 == False):
                    if ("empilhadeira" in agente) or ("transpaleteira" in agente):
                        prog_NR11 = True

                # testa se é necessário o programa NR12
                if (grupo == "acidente") and (prog_NR12 == False):
                    if ("operação de máquinas e equipamentos" in agente):
                        prog_NR12 = True

                # testa se é necessário o programa NR13
                if (grupo == "acidente") and (prog_NR13 == False):
                    if ("vaso de pressão" in agente):
                        prog_NR13 = True

                # testa se é necessário o programa NR33
                if (grupo == "acidente") and (prog_NR33 == False):
                    if ("espaço confinado" in agente):
                        prog_NR33 = True


                # testa se é necessário o programa NR35
                if (grupo == "acidente") and (prog_NR35 == False):
                    if ("trabalho em altura" in agente):
                        prog_NR35 = True

        dados_result = {
            "prog_auditivo": prog_auditivo,
            "prog_respiratorio": prog_respiratorio,
            "prog_NR6": prog_NR6,
            "prog_NR10": prog_NR10,
            "prog_NR11": prog_NR11,
            "prog_NR12": prog_NR11,
            "prog_NR13": prog_NR13,
            "prog_NR33": prog_NR33,
            "prog_NR35": prog_NR35,
            "nome_empresa": nome_empresa,
            "cnpj": cnpj,
            "cnae": cnae,
            "descr_cnae": cnae_desc,
            "grau_risco": dados_extraidos["Grau de Risco"],
            "masculino": masculino,
            "feminino": feminino,
            "menor": menor,
            "total": total
            }
        
        return dados_result
    

    def ler_tabelas_metodo2(self, tables):
        prog_auditivo = False
        prog_respiratorio = False
        prog_NR6 = False
        prog_NR10 = False
        prog_NR11 = False
        prog_NR12 = False
        prog_NR13 = False
        prog_NR33 = False
        prog_NR35 = False
        print("Processando regras de negócio...")

        for n in range(len(tables)):
            # tables[num_tabela][coluna][linha]

            if str(tables[n][0][0]).strip() == "Identificação":
                nome_empresa = str(tables[n+1][0][0]).replace('Empresa ','').strip()
                cnpj = str(tables[n+2][2][0]).replace('CNPJ ','').strip()
                cnae = str(tables[n+4][0][0]).replace('CNAE ','').strip()
                descr_cnae = str(tables[n+4][2][0]).replace('Descrição CNAE ','').strip()
                grau_risco = str(tables[n+4][1][0]).replace('Grau de Risco ','').strip()

            elif str(tables[n][0][0]).strip() == "Total de Funcionários": 
                masculino = str(tables[n+3][0][0]).strip()
                feminino = str(tables[n+3][1][0]).strip()
                menor = str(tables[n+3][2][0]).strip()
                total = str(tables[n+3][3][0]).strip()
                
            elif str(tables[n][0][0]).strip() == 'Agente': 
                agente = str(tables[n][1][0]).lower()
                grupo = str(tables[n][3][0]).lower().strip()
                idx = n

                while True:
                    idx +=1
                    if str(tables[idx][0][0]).strip() == "Orientação":
                        descricao_epis = str(tables[idx+1][0][0]).strip()
                        break
                    elif "EPI's" in str(tables[idx][0][0]).strip():
                        descricao_epis = str(tables[idx][0][0]).strip()
                        break
                    elif str(tables[idx][0][0]).strip() == 'Agente':
                        descricao_epis = "Não aplicado"
                        break
                    elif str(tables[idx][0][0]).strip() == "Total de Funcionários":
                        descricao_epis = "Não aplicado"
                        break

                # testa se é necessário o programa auditivo
                if ("RUÍDO CONTINUO (ACIMA DE 85 dB)".lower() in agente) and (prog_auditivo == False): 
                    prog_auditivo = True   
                elif ("RUÍDO CONTINUO (RUÍDOS ENTRE 80 A 85 dB)".lower() in agente) and (prog_auditivo == False): 
                    prog_auditivo = True
                    
                # testa se é necessário o programa respiratorio
                risco_respiratorio = ["químico", "biológico"]
                if (grupo in risco_respiratorio) and (prog_respiratorio == False): 
                    prog_respiratorio = True

                # testa se é necessário o programa NR6
                risco_nr6 = ["químico", "físico", "acidente", "biológico"]
                if (grupo in risco_nr6) and (prog_NR6 == False): 
                    lista_dispensa_de_epi = ["Não há necessidade", "Não aplicado", "Demais EPI's estão dispensados"]
                    for dispensa_de_epi in lista_dispensa_de_epi:
                        if dispensa_de_epi not in descricao_epis:
                            prog_NR6 = True

                # testa se é necessário o programa NR10
                if (grupo == "acidente") and (prog_NR10 == False):
                    if ("arco elétrico" in agente):
                        prog_NR10 = True

                # testa se é necessário o programa NR11
                if (grupo == "acidente") and (prog_NR11 == False):
                    if ("empilhadeira" in agente) or ("transpaleteira" in agente):
                        prog_NR11 = True

                # testa se é necessário o programa NR12
                if (grupo == "acidente") and (prog_NR12 == False):
                    if ("operação de máquinas e equipamentos" in agente):
                        prog_NR12 = True

                # testa se é necessário o programa NR13
                if (grupo == "acidente") and (prog_NR13 == False):
                    if ("vaso de pressão" in agente):
                        prog_NR13 = True

                # testa se é necessário o programa NR33
                if (grupo == "acidente") and (prog_NR33 == False):
                    if ("espaço confinado" in agente):
                        prog_NR33 = True


                # testa se é necessário o programa NR35
                if (grupo == "acidente") and (prog_NR35 == False):
                    if ("trabalho em altura" in agente):
                        prog_NR35 = True

        dados_result = {
            "prog_auditivo": prog_auditivo,
            "prog_respiratorio": prog_respiratorio,
            "prog_NR6": prog_NR6,
            "prog_NR10": prog_NR10,
            "prog_NR11": prog_NR11,
            "prog_NR12": prog_NR11,
            "prog_NR13": prog_NR13,
            "prog_NR33": prog_NR33,
            "prog_NR35": prog_NR35,
            "nome_empresa": nome_empresa,
            "cnpj": cnpj,
            "cnae": cnae,
            "descr_cnae": descr_cnae,
            "grau_risco": grau_risco,
            "masculino": masculino,
            "feminino": feminino,
            "menor": menor,
            "total": total
            }
        
        return dados_result


    def __init__(self, file_base_rtf:str):
        kill_process_word()
        time.sleep(1)
        self.file_base_rtf = file_base_rtf
        self.tables = self.doc_rtf_to_dataframe(file_base_rtf)
        self.data_vigencia = self.calc_data_vigencia(file_base_rtf)
        self.data_elaboracao = datetime.now().strftime('%d.%m.%Y')
        self.data_conclusao = datetime.now().strftime('%d %B DE %Y').upper()
        # programas de conservacao
        try:
            self.dados_result = self.ler_tabelas_metodo1(self.tables, self.file_base_rtf)
        except:
            try:
                self.dados_result = self.ler_tabelas_metodo2(self.tables)
            except:
                raise Exception("Erro, Nao foi possivel organizar os dados do documento de entrada, possivelmente a estrutura ficou diferente ao transformar em HTML")

        # Dados da Empresa
        self.nome_empresa = str(self.dados_result['nome_empresa']).replace('Empresa ','').strip()
        self.cnpj = str(self.dados_result['cnpj']).replace('CNPJ ','').strip()
        self.cnae = str(self.dados_result['cnae']).replace('CNAE ','').strip()
        self.descr_cnae = str(self.dados_result['descr_cnae']).replace('Descrição CNAE ','').strip()

        # Grau de Risco
        self.grau_risco = str(self.dados_result['grau_risco']).replace('Grau de Risco ','').strip()
        if self.grau_risco == '1':
            self.gr1, self.gr2, self.gr3, self.gr4 = '1', 'g.risc', 'g.risc', 'g.risc'
        if self.grau_risco == '2':
            self.gr1, self.gr2, self.gr3, self.gr4 = 'g.risc', '2', 'g.risc', 'g.risc'
        if self.grau_risco == '3':
            self.gr1, self.gr2, self.gr3, self.gr4 = 'g.risc', 'g.risc', '3', 'g.risc'
        if self.grau_risco == '4':
            self.gr1, self.gr2, self.gr3, self.gr4 = 'g.risc', 'g.risc', 'g.risc', '4'


        # Qtd Funcionarios
        self.masculino = str(self.dados_result['masculino']).strip()
        self.feminino = str(self.dados_result['feminino']).strip()
        self.menor = str(self.dados_result['menor']).strip()
        self.total = str(self.dados_result['total']).strip()

        # Referencia para pintar celular qtd funcionarios CIPA
        tt = int(self.total)
        if tt <= 19: self.ref_cipa = "0 a 19"
        elif tt >  19  and tt <  30:  self.ref_cipa = "20 a 29"
        elif tt >  29  and tt <  51:  self.ref_cipa = "30 a 50"
        elif tt >  50  and tt <  81:  self.ref_cipa = "51 a 80"
        elif tt >  80  and tt < 101:  self.ref_cipa = "81 a 100"
        elif tt > 100  and tt < 121:  self.ref_cipa = "101 a 120"
        elif tt > 120  and tt < 141:  self.ref_cipa = "121 a 140"
        elif tt > 140  and tt < 301:  self.ref_cipa = "141 a 300"
        elif tt > 300  and tt < 501:  self.ref_cipa = "301 a 500"
        elif tt > 500  and tt < 1001: self.ref_cipa = "501 a 1000"
        elif tt > 1000 and tt < 2501: self.ref_cipa = "1001 a 2500"
        elif tt > 2500 and tt < 5001: self.ref_cipa = "2501 a 5000"
        elif tt > 5000 and tt <= 10000: self.ref_cipa = "5001 a 10.000"
        elif tt > 10000: self.ref_cipa = "Acima de 10.000"

        # Programas de conservação
        self.prog_auditivo = bool(self.dados_result['prog_auditivo'])
        self.prog_respiratorio = bool(self.dados_result['prog_respiratorio'])
        self.prog_NR6 = bool(self.dados_result['prog_NR6'])
        self.prog_NR10 = bool(self.dados_result['prog_NR10'])
        self.prog_NR11 = bool(self.dados_result['prog_NR11'])
        self.prog_NR12 = bool(self.dados_result['prog_NR12'])
        self.prog_NR13 = bool(self.dados_result['prog_NR13'])
        self.prog_NR33 = bool(self.dados_result['prog_NR33'])
        self.prog_NR35 = bool(self.dados_result['prog_NR35'])

        # Numero de Integrantes necessários CIPA


        self.keys_rtf = {
            "ref_nome_empresa": self.nome_empresa,
            "ref_data_vigencia": self.data_vigencia,
            "ref_campo_data": self.data_elaboracao,
            "ref_data_conclusao": self.data_conclusao,
            "ref_cnae": self.cnae,
            "ref_grau_risco": self.grau_risco,
            "ref_descr_cnae": self.descr_cnae,
            "ref_masculino": self.masculino,
            "ref_feminino": self.feminino,
            "ref_menor": self.menor,
            "ref_total": self.total,
            "g.risc1": self.gr1,
            "g.risc2": self.gr2,
            "g.risc3": self.gr3,
            "g.risc4": self.gr4
        }


    
